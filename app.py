import pathlib
import time
import re
from tqdm import tqdm
from google import genai
from google.genai import types
from googletrans import Translator
from rich.console import Console
from rich.markdown import Markdown
from rich.panel import Panel
from rich.text import Text
from docx import Document
import typer
import pyfiglet

GENAI_API_KEY = ""
console = Console()

def trans(text):
    translator = Translator()
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    translated_paragraphs = []
    
    red_desc = "\033[31mTranslating\033[0m"  # tqdm desc에 ANSI 컬러 코드 넣은 예시
    for p in tqdm(paragraphs, desc=red_desc, unit="paragraph"):
        try:
            translated = translator.translate(p, src='en', dest='ko').text
            translated_paragraphs.append(translated)
        except Exception as e:
            translated_paragraphs.append(f"[Translation error: {e}]")
    return '\n\n'.join(translated_paragraphs)

def clean_text_for_markdown(text):
    text = re.sub(r'\n[-*]{3,}\n', '\n\n---\n\n', text)
    text = re.sub(r'\n(\* +)', r'\n\n\1', text)
    paragraphs = text.split('\n\n')
    cleaned = '\n\n'.join([p.strip() for p in paragraphs if p.strip()])
    return cleaned

def save_summary_to_docx(text, original_pdf: pathlib.Path):
    """요약 내용을 DOCX로 저장"""
    doc = Document()
    doc.add_heading(f"{original_pdf.stem} - 요약본", level=1)
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph)
    output_path = original_pdf.with_name(f"{original_pdf.stem}_summary.docx")
    doc.save(output_path)
    return output_path

def print_welcome():
    ascii_art = pyfiglet.figlet_format("WELCOME")
    console.print(Panel(Text(ascii_art, style="bold cyan"), border_style="cyan"))

def print_goodbye():
    ascii_art = pyfiglet.figlet_format("GOODBYE!")
    console.print(Panel(Text(ascii_art, style="bold magenta"), border_style="magenta"))

def main(pdf_path: pathlib.Path):
    """
    PDF 파일 경로를 받아서 요약 + 번역 후 출력 + DOCX 저장
    """
    print_welcome()
    console.print("[bold blue]프로그램 시작[/bold blue]")
    start_time = time.time()

    client = genai.Client(api_key=GENAI_API_KEY)
    console.print(f"[green]API 클라이언트 생성 완료[/green]")

    if not pdf_path.exists():
        console.print(f"[bold red]파일이 존재하지 않습니다:[/bold red] {pdf_path}")
        raise typer.Exit()
    console.print(f"[green]PDF 파일 경로 확인 완료:[/green] {pdf_path}")

    prompt = "Summarize this document"
    console.print("[green]요약 요청 시작[/green]")

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[
            types.Part.from_bytes(
                data=pdf_path.read_bytes(),
                mime_type='application/pdf',
            ),
            prompt
        ]
    )
    console.print("[green]요약 응답 받음[/green]")

    original_text = response.text

    try:
        console.print("[green]번역 시작[/green]")
        trans_result = trans(original_text)
        console.print("[green]번역 완료[/green]")

        cleaned_text = clean_text_for_markdown(trans_result)
        console.print("[green]마크다운 정리 완료[/green]")
        
        # console.print("\n" + "="*50)
        console.print("\n" + "="*50, style="yellow")
        console.print(Markdown(cleaned_text))

        # DOCX 저장
        # console.print("\n" + "="*50)
        console.print("\n" + "="*50, style="yellow")

        console.print("[green]DOCX 저장 시작[/green]")
        saved_path = save_summary_to_docx(cleaned_text, pdf_path)
        console.print(f"[bold red]요약본 저장 완료:[/bold red] {saved_path}")

    except Exception as e:
        console.print(f"[bold red]에러 발생:[/bold red] {e}")
        cleaned_text = clean_text_for_markdown(original_text)
        console.print(Markdown(cleaned_text))

    # console.print("\n" + "="*50)
    end_time = time.time()
    console.print(f"[bold blue]총 걸린 시간 : {end_time - start_time:.2f} 초[/bold blue]")
    print_goodbye()

if __name__ == "__main__":
    typer.run(main)

# 실행 예시:
# python3 app3.py input.pdf
